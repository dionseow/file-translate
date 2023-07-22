import os
from typing import Callable
import os
from docx import Document
from app.schemas.file import FileTranslator

class WordDocumentTranslator(FileTranslator):
    def translate_in_place(self, file_path: str | os.PathLike, new_file_path: str | os.PathLike, translator: Callable[[str], str]) -> None:
        doc = Document(file_path)

        # Translate paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text = paragraph.text
                translated_text = translator(text)
                paragraph.text = paragraph.text.replace(text, translated_text)
        
        # Translate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            text = paragraph.text
                            translated_text = translator(text)
                            paragraph.text = paragraph.text.replace(text, translated_text)

        # No support for graphs - https://python-docx.readthedocs.io/en/latest/user/shapes.html

        doc.save(new_file_path)