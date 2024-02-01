import os
from typing import Callable
import os
from docx import Document
from schemas.file import Processor

class WordDocumentProcessor(Processor):
    def extract_text(self, file_path: str | os.PathLike) -> str:
        doc = Document(file_path)

        extracted_text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text = paragraph.text
                extracted_text = extracted_text + " /n " + text
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            text = paragraph.text
                            extracted_text = extracted_text + " /n " + text
        
        return extracted_text

    def translate_in_place(self, file_path: str | os.PathLike, new_file_path: str | os.PathLike, translator: Callable[[str], str]) -> None:
        pass
        # doc = Document(file_path)

        # # Translate paragraphs
        # for paragraph in doc.paragraphs:
        #     if paragraph.text:
        #         text = paragraph.text
        #         translated_text = translator(text)
        #         paragraph.text = paragraph.text.replace(text, translated_text)
        
        # # Translate tables
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             for paragraph in cell.paragraphs:
        #                 if paragraph.text:
        #                     text = paragraph.text
        #                     translated_text = translator(text)
        #                     paragraph.text = paragraph.text.replace(text, translated_text)

        # # No support for graphs - https://python-docx.readthedocs.io/en/latest/user/shapes.html

        # doc.save(new_file_path)